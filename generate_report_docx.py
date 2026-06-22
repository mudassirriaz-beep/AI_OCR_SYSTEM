"""
Generate professional DOCX report for OCR System project.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

OUTPUT = r"C:\Users\ZAH\Downloads\OCR_System_Report.docx"
LOGO   = r"C:\Users\ZAH\Pictures\logo.png"

# Colours
C_DARK_TEAL  = RGBColor(0x0d, 0x4f, 0x5c)
C_TEAL       = RGBColor(0x1a, 0x7a, 0x8a)
C_ACCENT     = RGBColor(0x2e, 0xcc, 0x71)
C_LIGHT_TEAL = RGBColor(0xe8, 0xf4, 0xf6)
C_DARK_GREY  = RGBColor(0x2c, 0x3e, 0x50)
C_MID_GREY   = RGBColor(0x7f, 0x8c, 0x8d)
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT_GREY = RGBColor(0xf5, 0xf6, 0xfa)
C_GREEN_ROW  = RGBColor(0xe8, 0xf8, 0xee)


def hex_to_rgb_str(r, g, b):
    return f"{r:02X}{g:02X}{b:02X}"


DARK_TEAL_HEX  = "0d4f5c"
TEAL_HEX       = "1a7a8a"
LIGHT_TEAL_HEX = "e8f4f6"
LIGHT_GREY_HEX = "f5f6fa"
GREEN_ROW_HEX  = "e8f8ee"
WHITE_HEX      = "FFFFFF"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val', 'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz', '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'cde8ec'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_para(doc, text, bold=False, size=10, color=C_DARK_GREY,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4,
             italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    return p


def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    # Background shading via paragraph border
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  LIGHT_TEAL_HEX)
    pPr.append(shd)
    # Bottom border
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), TEAL_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = C_DARK_TEAL
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = C_TEAL
    return p


def make_table(doc, data, widths_cm, header=True):
    rows = len(data)
    cols = len(data[0])
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'

    for i, row_data in enumerate(data):
        row = t.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]

            if header and i == 0:
                set_cell_bg(cell, DARK_TEAL_HEX)
                run = p.add_run(str(cell_text))
                run.bold = True
                run.font.size  = Pt(9)
                run.font.color.rgb = C_WHITE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                bg = LIGHT_GREY_HEX if i % 2 == 0 else WHITE_HEX
                set_cell_bg(cell, bg)
                run = p.add_run(str(cell_text))
                run.font.size  = Pt(9)
                run.font.color.rgb = C_DARK_GREY
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            set_cell_border(cell)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)

    # Set column widths
    for i, row in enumerate(t.rows):
        for j, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[j])

    return t


def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ═══════════════════════════════════════════
    #  PAGE 1 — COVER
    # ═══════════════════════════════════════════

    # Logo
    if os.path.exists(LOGO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run()
        run.add_picture(LOGO, width=Cm(5))

    # Company name
    add_para(doc, "Al-Khair Institute of Technology",
             bold=True, size=12, color=C_DARK_TEAL,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # Title block with dark teal background
    title_tbl = doc.add_table(rows=3, cols=1)
    title_tbl.style = 'Table Grid'
    titles = [
        ("OCR System", True, 28, C_WHITE),
        ("AI-Powered Document Extraction Platform", False, 13, RGBColor(0xd0, 0xf0, 0xf5)),
        ("CNIC & Driving License — Automated Data Extraction", False, 10, RGBColor(0xa8, 0xd8, 0xe0)),
    ]
    for i, (text, bold, size, color) in enumerate(titles):
        cell = title_tbl.rows[i].cells[0]
        set_cell_bg(cell, DARK_TEAL_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8 if i == 0 else 2)
        p.paragraph_format.space_after  = Pt(8 if i == 2 else 2)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color
        # Remove border
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top','left','bottom','right'):
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), 'none')
            tcBorders.append(tag)
        tcPr.append(tcBorders)

    doc.add_paragraph()

    # Meta table
    meta = [
        ["Organization",  "Al-Khair Institute of Technology (AIT)"],
        ["Product Name",  "OCR System"],
        ["Document Type", "Internal Technical Report"],
        ["Version",       "V4.0 — Current Stable Release"],
        ["Report Date",   "June 19, 2026"],
        ["Status",        "Confidential — Internal Use Only"],
    ]
    make_table(doc, meta, [4.5, 12.5], header=False)

    doc.add_paragraph()
    add_h2(doc, "Project Team")
    team = [
        ["Name", "Role"],
        ["Muhammad Khurram Aimad", "Team Lead"],
        ["Syed Mudassir Ali",      "Main Developer"],
        ["Nehal Ansari",           "Assistant Developer"],
    ]
    make_table(doc, team, [8.5, 8.5])

    doc.add_page_break()

    # ═══════════════════════════════════════════
    #  PAGE 2 — OVERVIEW + ARCHITECTURE
    # ═══════════════════════════════════════════
    add_section_header(doc, "1.  Executive Summary & Project Overview")
    add_para(doc,
        "The OCR System is an AI-powered platform developed by Al-Khair Institute of Technology "
        "to automate data extraction from Pakistani identity documents (CNIC and Driving License). "
        "The system uses a fine-tuned Donut (Document Understanding Transformer) model to extract "
        "all key fields and automatically populate HTML forms — eliminating manual data entry, "
        "reducing errors, and accelerating document processing. Delivered as a standalone Windows "
        "desktop application, it requires no internet, no GPU, and no external installations.",
        size=10, color=C_DARK_GREY, space_after=6)

    highlights = [
        ["✓ 90.5% field-level accuracy (V4)", "✓ Supports CNIC and Driving License"],
        ["✓ Fills all HTML input types (text, date, number, email)", "✓ Extracts photograph from document"],
        ["✓ Fully offline — no internet required", "✓ Single 871 MB installer EXE"],
    ]
    make_table(doc, highlights, [8.5, 8.5], header=False)

    doc.add_paragraph()
    add_section_header(doc, "2.  System Architecture & Technology Stack")
    add_h2(doc, "Processing Pipeline")
    pipeline = [
        ["Step", "Component", "Description"],
        ["1–2", "Image Input",       "Accepts JPEG, PNG, WebP, PDF, HEIC. Auto-converts to JPEG."],
        ["3",   "Donut Inference",   "Fine-tuned transformer reads image → outputs structured JSON with all fields."],
        ["4",   "Photo Extraction",  "OpenCV face detection crops the photograph. Fallback to fixed top-right region."],
        ["5",   "Field Matching",    "RapidFuzz fuzzy matching maps extracted fields to HTML form inputs."],
        ["6",   "Value Formatting",  "Dates → ISO 8601, numbers → digits only, checkboxes → checked if truthy."],
        ["7",   "Form Output",       "Filled HTML form saved and viewable in browser or in-app preview panel."],
    ]
    make_table(doc, pipeline, [1.2, 3.5, 12.3])

    doc.add_paragraph()
    add_h2(doc, "Technology Stack")
    tech = [
        ["Category", "Technology", "Notes"],
        ["AI Model",     "Donut (naver-clova-ix/donut-base)",              "Fine-tuned on Pakistani CNIC/DL dataset"],
        ["Framework",    "PyTorch (CPU) + HuggingFace Transformers 4.46.3","CPU-only — no GPU required on client"],
        ["Image",        "OpenCV, Pillow, PyMuPDF",                        "Processing, face detection, PDF conversion"],
        ["Form Parsing", "BeautifulSoup4 + RapidFuzz",                     "HTML parsing, fuzzy field matching"],
        ["UI",           "Tkinter (Python built-in)",                      "Desktop GUI, no web server needed"],
        ["Packaging",    "PyInstaller + Inno Setup 6",                     "871 MB self-contained installer EXE"],
        ["Platform",     "Windows 10 / 11 (64-bit)",                       "Python 3.11"],
    ]
    make_table(doc, tech, [3.0, 7.0, 7.0])

    doc.add_page_break()

    # ═══════════════════════════════════════════
    #  PAGE 3 — AI MODEL + TRAINING
    # ═══════════════════════════════════════════
    add_section_header(doc, "3.  AI Model — Donut")
    add_para(doc,
        "Donut (Document Understanding Transformer) is an end-to-end transformer model that reads "
        "a document image and outputs structured text with no traditional OCR engine. It uses a "
        "Swin Transformer image encoder and a BART-based decoder to generate structured JSON directly "
        "from the image — making it robust to poor image quality, worn cards, and varying camera angles. "
        "The fine-tuned model checkpoint is 777 MB, bundled inside the installer.",
        size=10, color=C_DARK_GREY, space_after=6)

    add_h2(doc, "Why Donut was chosen over traditional OCR + SLM pipeline")
    reasons = [
        ["End-to-end learning — reads and parses in one step, no separate OCR engine",
         "Robust to poor lighting, camera angles, worn cards"],
        ["Fine-tunable on custom data (Pakistani CNIC/DL)",
         "CPU inference — no GPU required on client"],
        ["Replaces entire OCR + rules + SLM chain with a single model",
         "Structured JSON output — no regex post-processing needed"],
    ]
    make_table(doc, reasons, [8.5, 8.5], header=False)

    doc.add_paragraph()
    add_section_header(doc, "4.  Training History & Accuracy")
    add_para(doc,
        "The model was fine-tuned in four iterative rounds on a dataset of 1,177 human-reviewed "
        "CNIC and Driving License images (1,001 train / 176 validation). Each round built on the "
        "previous checkpoint. Training used an NVIDIA GeForce RTX 5050 GPU.",
        size=10, color=C_DARK_GREY, space_after=6)

    train = [
        ["Round", "Version", "Accuracy", "Key Changes"],
        ["1", "V1", "0%",     "Model failed to converge — task token format was incorrect."],
        ["2", "V2", "81.5%",  "Fixed task token format. Model successfully learned document structure."],
        ["3", "V3", "89.4%",  "Expanded dataset with diverse samples. Improved date field accuracy."],
        ["4", "V4 ★","90.5%", "Further data expansion. Best checkpoint Epoch 1 (val_loss=0.0373). Current stable release."],
    ]
    t = make_table(doc, train, [1.5, 2.0, 2.5, 11.0])
    # Highlight V4 row green
    for j in range(4):
        set_cell_bg(t.rows[4].cells[j], GREEN_ROW_HEX)

    add_para(doc, "★  V4 is the current stable release shipped to clients.",
             italic=True, size=8, color=C_MID_GREY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    #  PAGE 4 — FEATURES + RESULTS
    # ═══════════════════════════════════════════
    add_section_header(doc, "5.  Features & Capabilities")
    features = [
        ["Area", "Capability"],
        ["Document Support",  "Pakistani CNIC and Driving License  |  JPEG, PNG, WebP, BMP, TIFF, HEIC, GIF, PDF"],
        ["Extracted Fields",  "Full Name, Father Name, Gender, ID/License Number, DOB, Date of Issue, Date of Expiry, Country, Province, Vehicle Category, Photograph"],
        ["Form Filling",      "Fills any HTML form — text, date (ISO conversion), number, email, tel, select, radio, checkbox, textarea  |  Fuzzy field-name matching"],
        ["UI Features",       "Real-time processing log  |  Extracted fields sidebar  |  'View Extracted Data' in-app popup (photo + all fields)  |  Open filled form in browser"],
        ["Deployment",        "Fully offline  |  No GPU required  |  No Python/Ollama installs  |  Single 871 MB EXE installer  |  Windows 10/11 (64-bit)"],
    ]
    make_table(doc, features, [3.5, 13.5])

    doc.add_paragraph()
    add_section_header(doc, "6.  Results & Performance")
    results = [
        ["Metric", "Value"],
        ["Overall field accuracy (V4)",          "90.5%"],
        ["Total labeled dataset",                "1,177 documents  (1,001 train / 176 validation)"],
        ["Confidence threshold for review flag", "Below 70% → flagged for human verification"],
        ["Best validation loss (V4)",            "0.0373  (Epoch 1 of 7)"],
        ["Inference speed — CPU only",           "8–15 seconds per document"],
        ["Inference speed — GPU",                "2–4 seconds per document"],
        ["Installer size",                       "871 MB  (self-contained, no internet required)"],
    ]
    make_table(doc, results, [6.0, 11.0])

    doc.add_page_break()

    # ═══════════════════════════════════════════
    #  PAGE 5 — TEAM + CONCLUSION
    # ═══════════════════════════════════════════
    add_section_header(doc, "7.  Team Members & Roles")
    team_detail = [
        ["Name", "Role", "Responsibilities"],
        ["Muhammad Khurram Aimad", "Team Lead",
         "Project leadership, technical direction, architecture decisions, AI model strategy, fine-tuning roadmap, client requirements, quality assurance."],
        ["Syed Mudassir Ali", "Main Developer",
         "Donut model integration and fine-tuning, training scripts, form-filling engine, desktop application, PyInstaller packaging, Inno Setup installer."],
        ["Nehal Ansari", "Assistant Developer",
         "Dataset preparation, image labeling and review, testing and validation, UI enhancements, bug verification, and documentation support."],
    ]
    make_table(doc, team_detail, [4.0, 3.5, 9.5])

    doc.add_paragraph()
    add_section_header(doc, "8.  Conclusion & Future Work")
    add_para(doc,
        "The OCR System successfully delivers a production-ready AI document extraction solution "
        "achieving 90.5% field-level accuracy — improved from 0% in Round 1 to a robust stable "
        "release in Round 4 through iterative fine-tuning. The product is delivered as a self-contained "
        "871 MB Windows installer, immediately deployable in banks, hospitals, HR departments, and "
        "government offices without any client-side setup.",
        size=10, color=C_DARK_GREY, space_after=6)

    add_h2(doc, "Future Roadmap")
    future = [
        ["Further fine-tuning (2,000–3,000 samples) — target 94–96% accuracy",
         "Expansion to Passport, Birth Certificate, Utility Bills"],
        ["SaaS web platform with user accounts, billing, and REST API",
         "Mobile application for on-device document capture"],
    ]
    make_table(doc, future, [8.5, 8.5], header=False)

    doc.add_paragraph()
    # Footer note
    end_tbl = doc.add_table(rows=1, cols=1)
    cell = end_tbl.rows[0].cells[0]
    set_cell_bg(cell, DARK_TEAL_HEX)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(
        "Al-Khair Institute of Technology (AIT)  ·  OCR System v4.0  ·  "
        "Internal Technical Report  ·  June 2026  ·  Confidential")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = C_WHITE

    doc.save(OUTPUT)
    print(f"DOCX saved: {OUTPUT}")


if __name__ == "__main__":
    build()
